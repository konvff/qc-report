"""
Core report-generation engine (v2).

KEY FIX: the master template is a REAL completed report (Bjorna / Coca-Cola PO
13320), not a blank form. Every checkbox and photo in it that we don't
explicitly overwrite would otherwise leak into every new report. This version:

1. Runs a global sanitizer on a fresh in-memory copy of the template before
   filling anything: force-unchecks every legacy FORMCHECKBOX field in the
   whole document, clears the one manually-inserted Wingdings tick (used for
   the Inspection Conclusion boxes), and strips every photo out of the 6
   photo-gallery tables. This guarantees a genuinely blank starting point
   every single time, regardless of whether a given field is wired up yet.

2. Handles checkboxes via the REAL mechanism Word uses in this file: legacy
   FORMCHECKBOX form fields (<w:checkBox> with <w:default>/<w:checked>), not
   plain text. A single cell can contain more than one checkbox (e.g. a
   "YES [ ] NO [ ]" pair), so checkboxes are located and disambiguated by the
   label text that immediately follows them.

3. Derives photo slots (row, col, default caption) directly from the
   template at import time, so nothing about section layout is hardcoded by
   hand and titles can be freely edited by the user -- the generator just
   places whatever title text + image it's given at that exact slot.
"""
import io
import copy as _copy
from docx import Document
from docx.shared import Emu, Pt
from docx.oxml.ns import qn

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

PHOTO_TABLE_INDEXES = {
    "standards_photos": 7,
    "packing_photos": 10,
    "defect_photos": 15,
    "measurement_photos": 17,
    "presentation_photos": 19,
    "shrinkage_photos": 20,
}

TABLES = {
    "header_info": 0,
    "product_category": 1,
    "po_details": 2,
    "aql_results": 3,
    "conclusion": 4,
    "po_comments": 5,
    "standards_reference": 6,
    "standards_photos": 7,
    "lab_test": 8,
    "packing_matrix": 9,
    "packing_photos": 10,
    "marking_labeling": 11,
    "upc_verification": 12,
    "cartons_selected": 13,
    "defects_log": 14,
    "defect_photos": 15,
    "measurement_chart": 16,
    "measurement_photos": 17,
    "onsite_tests": 18,
    "presentation_photos": 19,
    "shrinkage_photos": 20,
    "shrinkage_chart": 21,
}

DEFECT_TAXONOMY = [
    "Uneven Stitch", "Open Seam", "Color Stain", "Off Registration", "Pleat",
    "Weaving Defect", "Hanging Thread", "Shade Within Set", "Stain",
    "Uncut Thread", "Puckering", "Touching",
]


# ---------------------------------------------------------------------------
# Low-level cell text helpers
# ---------------------------------------------------------------------------
def _first_run_format(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            if r.text.strip() or r.font.size or r.font.name:
                return {
                    "name": r.font.name,
                    "size": r.font.size,
                    "bold": r.font.bold,
                    "italic": r.font.italic,
                    "color": r.font.color.rgb if r.font.color and r.font.color.type else None,
                }
    return {"name": "Cambria", "size": Emu(114300), "bold": None, "italic": None, "color": None}


def set_cell_text(cell, text, fmt=None):
    if fmt is None:
        fmt = _first_run_format(cell)
    p = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    run = p.add_run(str(text) if text is not None else "")
    run.font.name = fmt.get("name") or "Cambria"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), fmt.get("name") or "Cambria")
    if fmt.get("size"):
        run.font.size = fmt["size"]
    if fmt.get("bold") is not None:
        run.font.bold = fmt["bold"]
    if fmt.get("italic") is not None:
        run.font.italic = fmt["italic"]
    if fmt.get("color"):
        run.font.color.rgb = fmt["color"]
    return run


def insert_image_in_cell(cell, image_source, width_emu=2870000):
    """image_source: filesystem path. Normalizes/converts via Pillow when
    available so odd formats (HEIC, huge phone photos, CMYK JPEGs) don't
    crash python-docx. Returns True on success, False if the image couldn't
    be read at all (never raises -- one bad photo shouldn't kill the whole
    report generation)."""
    p = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    run = p.add_run()

    try:
        if PIL_AVAILABLE:
            img = Image.open(image_source)
            img = img.convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=88)
            buf.seek(0)
            run.add_picture(buf, width=Emu(width_emu))
        else:
            run.add_picture(image_source, width=Emu(width_emu))
        return True
    except Exception:
        return False


# ---------------------------------------------------------------------------
# Checkbox handling (legacy FORMCHECKBOX fields)
# ---------------------------------------------------------------------------
def _find_checkboxes_in_cell(cell):
    """Returns [(checkBox_element, label_text), ...] in document order. A
    single cell can hold more than one checkbox (e.g. 'YES [ ] NO [ ]')."""
    tc = cell._tc
    runs = list(tc.iter(f"{W_NS}r"))
    results = []
    for i, r in enumerate(runs):
        cb = r.find(f".//{W_NS}checkBox")
        if cb is None:
            continue
        label = ""
        for r2 in runs[i + 1:]:
            if r2.find(f".//{W_NS}checkBox") is not None:
                break
            t = r2.find(f"{W_NS}t")
            if t is not None and t.text:
                label += t.text
        results.append((cb, label.strip()))
    return results


def _set_checkbox_element(cb_el, checked: bool):
    for tag in ("default", "checked"):
        el = cb_el.find(f"{W_NS}{tag}")
        if el is not None:
            cb_el.remove(el)
    val = "1" if checked else "0"
    default_el = cb_el.makeelement(f"{W_NS}default", {f"{W_NS}val": val})
    cb_el.append(default_el)
    checked_el = cb_el.makeelement(f"{W_NS}checked", {f"{W_NS}val": val})
    cb_el.append(checked_el)


def resize_row_block(table, first_row_idx, current_block_len, new_len):
    """Clones or trims table rows so a repeatable block (e.g. PO rows, AQL
    rows) has exactly `new_len` rows, instead of leaving stale template rows
    behind or silently dropping data that doesn't fit. Returns the number of
    rows the block now has (>=1, a block is never fully removed)."""
    new_len = max(1, new_len)
    tbl = table._tbl
    trs = tbl.findall(f"{W_NS}tr")
    block_trs = trs[first_row_idx: first_row_idx + current_block_len]
    if not block_trs:
        return current_block_len
    if new_len <= current_block_len:
        for tr in block_trs[new_len:]:
            tbl.remove(tr)
    else:
        insert_after = block_trs[-1]
        for _ in range(new_len - current_block_len):
            new_tr = _copy.deepcopy(insert_after)
            insert_after.addnext(new_tr)
            insert_after = new_tr
    return new_len


def check_cell(cell, label_contains=None):
    """Marks a checkbox as checked. Sets the underlying FORMCHECKBOX state
    (correct per spec, and what real Word reads) AND additionally inserts a
    visible Wingdings tick mark right next to the target checkbox's label --
    this second step is what actually guarantees the checkmark is visible
    when the file is opened in LibreOffice too, since LO does not reliably
    re-render a FORMCHECKBOX's checked state from a python-docx save."""
    boxes = _find_checkboxes_in_cell(cell)
    if not boxes:
        return False

    target = None
    if label_contains:
        for cb, label in boxes:
            if label_contains.lower() in label.lower():
                target = cb
                break
    else:
        target = boxes[0][0]
    if target is None:
        return False

    _set_checkbox_element(target, True)
    _insert_tick_after_checkbox(target)
    return True


def _insert_tick_after_checkbox(cb_element):
    """Inserts a visible Wingdings-2 tick-mark run immediately after the
    'end' fldChar of the given checkbox field, so it appears right before
    that checkbox's label text (e.g. '[tick]YES  [ ]NO')."""
    begin_run = cb_element.getparent().getparent().getparent()  # checkBox -> ffData -> fldChar -> r
    # walk forward through sibling runs to find this field's terminating 'end'
    sib = begin_run.getnext()
    end_run = None
    while sib is not None:
        if sib.tag == f"{W_NS}r":
            fld = sib.find(f"{W_NS}fldChar")
            if fld is not None and fld.get(f"{W_NS}fldCharType") == "end":
                end_run = sib
                break
        sib = sib.getnext()
    if end_run is None:
        return

    tick_run = begin_run.makeelement(f"{W_NS}r", {})
    rPr = tick_run.makeelement(f"{W_NS}rPr", {})
    rFonts = rPr.makeelement(f"{W_NS}rFonts", {f"{W_NS}ascii": "Wingdings 2", f"{W_NS}hAnsi": "Wingdings 2"})
    b = rPr.makeelement(f"{W_NS}b", {})
    sz = rPr.makeelement(f"{W_NS}sz", {f"{W_NS}val": "18"})
    rPr.append(rFonts)
    rPr.append(b)
    rPr.append(sz)
    tick_run.append(rPr)
    sym = tick_run.makeelement(f"{W_NS}sym", {f"{W_NS}font": "Wingdings 2", f"{W_NS}char": "F050"})
    tick_run.append(sym)
    end_run.addnext(tick_run)


# ---------------------------------------------------------------------------
# Global sanitizer -- ALWAYS run before filling, on an in-memory copy only.
# ---------------------------------------------------------------------------
def sanitize_document(doc):
    body = doc.element.body
    for cb in body.iter(f"{W_NS}checkBox"):
        _set_checkbox_element(cb, False)
    for sym in list(body.iter(f"{W_NS}sym")):
        font = sym.get(f"{W_NS}font") or ""
        if "wingdings" in font.lower():
            run = sym.getparent()
            parent = run.getparent()
            if parent is not None:
                parent.remove(run)
    # Strip every image from the document BODY (all tables, not just the 6
    # designated photo galleries -- the real report had ad-hoc reference
    # photos pasted into other cells too, e.g. onsite-test remarks). The
    # page-header/footer logo lives outside doc.element.body, so it's safe.
    for drawing in list(body.iter(f"{W_NS}drawing")):
        run = drawing.getparent()
        parent = run.getparent()
        if parent is not None:
            parent.remove(run)


# ---------------------------------------------------------------------------
# Dynamic photo-slot discovery -- read straight from the template so nothing
# is hand-transcribed and out of sync.
# ---------------------------------------------------------------------------
def discover_photo_slots(template_path):
    doc = Document(template_path)
    slots = {}
    for section, idx in PHOTO_TABLE_INDEXES.items():
        t = doc.tables[idx]
        ncols = len(t.columns)
        section_slots = []
        for img_row in range(1, len(t.rows), 2):
            cap_row = img_row + 1 if img_row + 1 < len(t.rows) else None
            row_ncols = min(ncols, len(t.rows[img_row].cells))
            for col in range(row_ncols):
                try:
                    caption = t.cell(cap_row, col).text.strip() if cap_row else ""
                except IndexError:
                    caption = ""
                section_slots.append({
                    "row": img_row,
                    "col": col,
                    "default_title": caption or f"Photo {len(section_slots) + 1}",
                })
        slots[section] = section_slots
    return slots


# ---------------------------------------------------------------------------
# Main generation
# ---------------------------------------------------------------------------
def generate_report(template_path, output_path, data):
    doc = Document(template_path)
    sanitize_document(doc)
    tables = doc.tables

    if "header_info" in data:
        h = data["header_info"]
        t = tables[TABLES["header_info"]]
        set_cell_text(t.cell(0, 1), h.get("report_no", ""))
        set_cell_text(t.cell(1, 1), h.get("customer_name", ""))
        set_cell_text(t.cell(1, 3), h.get("destination_country", ""))
        set_cell_text(t.cell(2, 1), h.get("inspection_type", ""))
        set_cell_text(t.cell(2, 3), h.get("inspection_date", ""))
        set_cell_text(t.cell(3, 1), h.get("manufacturer_name", ""))
        set_cell_text(t.cell(3, 3), h.get("inspection_location", ""))
        set_cell_text(t.cell(4, 1), h.get("factory_rep_name", ""))
        set_cell_text(t.cell(4, 3), h.get("arrival_time", ""))
        set_cell_text(t.cell(5, 1), h.get("inspector_names", ""))
        set_cell_text(t.cell(5, 3), h.get("departure_time", ""))

    if "product_category" in data:
        pc = data["product_category"]
        t = tables[TABLES["product_category"]]
        for i in range(1, 6):
            set_cell_text(t.cell(0, i), pc.get("category_description", ""))
        cat_col = {"BATH": 0, "KITCHEN": 1, "TABLE": 2, "BEDDING": 3, "WINDOW": 4, "OTHER": 5}
        chosen = pc.get("category")
        if chosen in cat_col:
            check_cell(t.cell(1, cat_col[chosen]))
        for i in (1, 2):
            set_cell_text(t.cell(2, i), pc.get("fabric_required", ""))
        for i in (4, 5):
            set_cell_text(t.cell(2, i), pc.get("fabric_found", ""))
        for i in (1, 2):
            set_cell_text(t.cell(3, i), pc.get("poly_required", ""))
        for i in (4, 5):
            set_cell_text(t.cell(3, i), pc.get("poly_found", ""))

    if "po_rows" in data:
        t = tables[TABLES["po_details"]]
        orig_data_rows = len(t.rows) - 3  # minus 2 header rows and 1 Total row
        n_data_rows = resize_row_block(t, 2, orig_data_rows, len(data["po_rows"]))
        t = doc.tables[TABLES["po_details"]]  # re-fetch after XML resize
        total_qty = 0
        for i in range(n_data_rows):
            r = 2 + i
            row_data = data["po_rows"][i] if i < len(data["po_rows"]) else {}
            set_cell_text(t.cell(r, 0), row_data.get("po_number", ""))
            set_cell_text(t.cell(r, 1), row_data.get("sku", ""))
            set_cell_text(t.cell(r, 2), row_data.get("item_description", ""))
            set_cell_text(t.cell(r, 3), row_data.get("design_color", ""))
            set_cell_text(t.cell(r, 4), row_data.get("size", ""))
            set_cell_text(t.cell(r, 5), row_data.get("pcs_set_ctn", ""))
            set_cell_text(t.cell(r, 6), row_data.get("po_qty", ""))
            set_cell_text(t.cell(r, 7), row_data.get("offer_qty_sets", ""))
            set_cell_text(t.cell(r, 8), row_data.get("offer_qty_carton", ""))
            set_cell_text(t.cell(r, 9), row_data.get("pct_vs_po", ""))
            try:
                total_qty += int(row_data.get("po_qty") or 0)
            except (ValueError, TypeError):
                pass
        total_row = len(t.rows) - 1
        total_carton = 0
        for i in range(n_data_rows):
            row_data = data["po_rows"][i] if i < len(data["po_rows"]) else {}
            try:
                total_carton += int(row_data.get("offer_qty_carton") or 0)
            except (ValueError, TypeError):
                pass
        set_cell_text(t.cell(total_row, 6), str(total_qty) if total_qty else "")
        set_cell_text(t.cell(total_row, 7), str(total_qty) if total_qty else "")
        set_cell_text(t.cell(total_row, 8), str(total_carton) if total_carton else "")
        set_cell_text(t.cell(total_row, 9), "")
        tables = doc.tables  # refresh full handle after structural change

    if "aql_rows" in data:
        t = tables[TABLES["aql_results"]]
        orig_data_rows = len(t.rows) - 2  # minus header row and Total row
        n_data_rows = resize_row_block(t, 1, orig_data_rows, len(data["aql_rows"]))
        t = doc.tables[TABLES["aql_results"]]
        sums = [0] * 6  # sample_size, critical_found, critical_allowed, major_found, major_allowed, minor_found... (we'll compute what we can)
        total_sample = total_cf = total_ca = total_mjf = total_mja = total_mnf = total_mna = 0
        for i in range(n_data_rows):
            r = 1 + i
            row_data = data["aql_rows"][i] if i < len(data["aql_rows"]) else {}
            set_cell_text(t.cell(r, 0), row_data.get("item_description", ""))
            set_cell_text(t.cell(r, 1), row_data.get("size", ""))
            set_cell_text(t.cell(r, 2), row_data.get("sample_size", ""))
            set_cell_text(t.cell(r, 3), row_data.get("critical_found", ""))
            set_cell_text(t.cell(r, 4), row_data.get("critical_allowed", ""))
            set_cell_text(t.cell(r, 5), row_data.get("major_found", ""))
            set_cell_text(t.cell(r, 6), row_data.get("major_allowed", ""))
            set_cell_text(t.cell(r, 7), row_data.get("minor_found", ""))
            set_cell_text(t.cell(r, 8), row_data.get("minor_allowed", ""))
            set_cell_text(t.cell(r, 9), row_data.get("pass_fail", ""))

            def _num(v):
                try:
                    return int(v)
                except (ValueError, TypeError):
                    return 0
            total_sample += _num(row_data.get("sample_size"))
            total_cf += _num(row_data.get("critical_found"))
            total_ca += _num(row_data.get("critical_allowed"))
            total_mjf += _num(row_data.get("major_found"))
            total_mja += _num(row_data.get("major_allowed"))
            total_mnf += _num(row_data.get("minor_found"))
            total_mna += _num(row_data.get("minor_allowed"))

        total_row = len(t.rows) - 1
        set_cell_text(t.cell(total_row, 2), str(total_sample) if total_sample else "")
        set_cell_text(t.cell(total_row, 3), f"{total_cf:02d}")
        set_cell_text(t.cell(total_row, 4), f"{total_ca:02d}")
        set_cell_text(t.cell(total_row, 5), str(total_mjf))
        set_cell_text(t.cell(total_row, 6), str(total_mja))
        set_cell_text(t.cell(total_row, 7), str(total_mnf))
        set_cell_text(t.cell(total_row, 8), str(total_mna))
        overall_pass = total_mjf <= total_mja and total_mnf <= total_mna and total_cf <= total_ca
        set_cell_text(t.cell(total_row, 9), "PASS" if overall_pass else "FAIL")
        tables = doc.tables

    if "defects" in data or "defects_meta" in data:
        t = tables[TABLES["defects_log"]]
        meta = data.get("defects_meta", {})
        set_cell_text(t.cell(0, 3), meta.get("product", ""))
        set_cell_text(t.cell(0, 4), meta.get("product", ""))
        set_cell_text(t.cell(1, 3), meta.get("size", ""))
        set_cell_text(t.cell(1, 4), meta.get("size", ""))
        set_cell_text(t.cell(2, 3), meta.get("sample_size", ""))
        set_cell_text(t.cell(2, 4), meta.get("sample_size", ""))
        set_cell_text(t.cell(3, 3), meta.get("color", ""))
        set_cell_text(t.cell(3, 4), meta.get("color", ""))

        defects = data.get("defects", {})
        total_major, total_minor = 0, 0
        for i, label in enumerate(DEFECT_TAXONOMY):
            r = 5 + i
            entry = defects.get(label, {})
            major = entry.get("major", "")
            minor = entry.get("minor", "")
            set_cell_text(t.cell(r, 0), str(i + 1))
            set_cell_text(t.cell(r, 1), label)
            set_cell_text(t.cell(r, 2), label)
            set_cell_text(t.cell(r, 3), major)
            set_cell_text(t.cell(r, 4), minor)
            try:
                total_major += int(major) if str(major).strip() else 0
            except ValueError:
                pass
            try:
                total_minor += int(minor) if str(minor).strip() else 0
            except ValueError:
                pass

        allowed_major = meta.get("major_allowed")
        allowed_minor = meta.get("minor_allowed")
        set_cell_text(t.cell(18, 3), f"{total_major:02d}")
        set_cell_text(t.cell(18, 4), f"{total_minor:02d}")
        if allowed_major not in (None, ""):
            set_cell_text(t.cell(19, 3), allowed_major)
        if allowed_minor not in (None, ""):
            set_cell_text(t.cell(19, 4), allowed_minor)
        try:
            result_major = "Pass" if allowed_major not in (None, "") and total_major <= int(allowed_major) else "Fail"
            result_minor = "Pass" if allowed_minor not in (None, "") and total_minor <= int(allowed_minor) else "Fail"
        except (ValueError, TypeError):
            result_major = result_minor = "Pending"
        set_cell_text(t.cell(20, 3), result_major)
        set_cell_text(t.cell(20, 4), result_minor)

    if "conclusion" in data:
        t = tables[TABLES["conclusion"]]
        col_map = {"CONFORM": 2, "NOT CONFORM": 5, "PENDING": 8}
        target_col = col_map.get(data["conclusion"])
        if target_col is not None:
            cell = t.cell(0, target_col)
            run = set_cell_text(cell, "")
            run.font.name = "Wingdings 2"
            run.font.size = Pt(16)
            run.font.bold = True
            sym = run._element.makeelement(f"{W_NS}sym", {f"{W_NS}font": "Wingdings 2", f"{W_NS}char": "F050"})
            run._element.append(sym)

    if "po_comments" in data:
        t = tables[TABLES["po_comments"]]
        orig_rows = len(t.rows) - 1
        n_rows = resize_row_block(t, 1, orig_rows, len(data["po_comments"]))
        t = doc.tables[TABLES["po_comments"]]
        for i in range(n_rows):
            r = 1 + i
            row_data = data["po_comments"][i] if i < len(data["po_comments"]) else {}
            set_cell_text(t.cell(r, 0), row_data.get("po_number", ""))
            set_cell_text(t.cell(r, 1), row_data.get("sku_style", ""))
            set_cell_text(t.cell(r, 2), row_data.get("color", ""))
            set_cell_text(t.cell(r, 3), row_data.get("comments", ""))
        tables = doc.tables

    if "standards_reference" in data:
        t = tables[TABLES["standards_reference"]]
        option_col = {
            "provided_office": 1, "provided_supplier": 2, "suppliers_counter": 3,
            "not_available": 4, "with_auth": 5, "without_auth": 6,
        }
        for key, row in (("reference_samples", 1), ("specification_file", 2)):
            choice = data["standards_reference"].get(key)
            if choice in option_col:
                check_cell(t.cell(row, option_col[choice]))

    if "lab_test" in data:
        t = tables[TABLES["lab_test"]]
        row_index = {"lab_test_exist": 1, "lab_report_reviewed": 2,
                     "lab_report_per_protocols": 3, "any_deviation": 4, "result": 5}
        mark_col = {"yes": 1, "no": 2, "na": 3, "equip_na": 4}
        for key, r in row_index.items():
            entry = data["lab_test"].get(key, {})
            mark = entry.get("mark")
            if mark in mark_col:
                check_cell(t.cell(r, mark_col[mark]))
            if entry.get("remark"):
                set_cell_text(t.cell(r, 5), entry["remark"])

    if "cartons_selected" in data:
        t = tables[TABLES["cartons_selected"]]
        orig_rows = len(t.rows) - 2
        n_rows = resize_row_block(t, 1, orig_rows, len(data["cartons_selected"]))
        t = doc.tables[TABLES["cartons_selected"]]
        for i in range(n_rows):
            r = 1 + i
            row_data = data["cartons_selected"][i] if i < len(data["cartons_selected"]) else {}
            set_cell_text(t.cell(r, 0), row_data.get("item_size", ""))
            set_cell_text(t.cell(r, 1), row_data.get("total_cartons", ""))
            set_cell_text(t.cell(r, 2), row_data.get("serial_range", ""))
            set_cell_text(t.cell(r, 3), row_data.get("num_selected", ""))
            set_cell_text(t.cell(r, 4), row_data.get("serial_selected", ""))
        tables = doc.tables

    if "upc_verification" in data:
        t = tables[TABLES["upc_verification"]]
        orig_rows = len(t.rows) - 1
        n_rows = resize_row_block(t, 1, orig_rows, len(data["upc_verification"]))
        t = doc.tables[TABLES["upc_verification"]]
        for i in range(n_rows):
            r = 1 + i
            row_data = data["upc_verification"][i] if i < len(data["upc_verification"]) else {}
            set_cell_text(t.cell(r, 0), row_data.get("po_number", ""))
            set_cell_text(t.cell(r, 1), row_data.get("item_description", ""))
            set_cell_text(t.cell(r, 2), row_data.get("size", ""))
            set_cell_text(t.cell(r, 3), row_data.get("upc_as_per_po", ""))
            set_cell_text(t.cell(r, 4), row_data.get("upc_observed", ""))
        tables = doc.tables

    if "measurements" in data:
        t = tables[TABLES["measurement_chart"]]
        # Delete existing rows except the header
        for row in t.rows[1:]:
            t._tbl.remove(row._tr)
            
        m_data = data["measurements"]
        if isinstance(m_data, list):
            last_desc_cell = None
            last_desc_text = None
            for item in m_data:
                if item.get("type") == "header":
                    r = t.add_row()
                    set_cell_text(r.cells[0], item.get("item_size", ""))
                    set_cell_text(r.cells[1], item.get("color", ""))
                    # Merge cells 1 through 12 for the header row
                    for i in range(2, min(13, len(r.cells))):
                        r.cells[1].merge(r.cells[i])
                    last_desc_cell = None
                    last_desc_text = None
                elif item.get("type") == "data":
                    r = t.add_row()
                    desc = item.get("desc", "")
                    set_cell_text(r.cells[0], desc)
                    set_cell_text(r.cells[1], item.get("point", ""))
                    set_cell_text(r.cells[2], item.get("spec", ""))
                    
                    for i in range(1, 11):
                        set_cell_text(r.cells[2 + i], item.get(f"c{i}", ""))
                        
                    # Handle vertical merge for Description column
                    if last_desc_text == desc and last_desc_cell is not None:
                        # Clear text in current cell before merging to avoid duplication
                        set_cell_text(r.cells[0], "")
                        last_desc_cell.merge(r.cells[0])
                    else:
                        last_desc_text = desc
                        last_desc_cell = r.cells[0]

    if "measurement_options" in data:
        opt = data["measurement_options"]
        for p in document.paragraphs:
            if "Buyer Measurement Chart" in p.text:
                cbs = p._p.xpath('.//w:ffData/w:checkBox')
                if len(cbs) >= 2:
                    _set_checkbox_element(cbs[0], opt.get("buyer_chart", False))
                    _set_checkbox_element(cbs[1], opt.get("supplier_chart", False))
            elif "Within Tolerance" in p.text:
                cbs = p._p.xpath('.//w:ffData/w:checkBox')
                if len(cbs) >= 3:
                    _set_checkbox_element(cbs[0], opt.get("within_tolerance", False))
                    _set_checkbox_element(cbs[1], opt.get("beyond_tolerance", False))
                    _set_checkbox_element(cbs[2], opt.get("actual_findings", False))

    if "onsite_tests" in data:
        t = tables[TABLES["onsite_tests"]]
        row_index = {"needle_detection": 1, "metal_detector": 2, "carton_drop_test": 3,
                     "gsm": 4, "barcode_scan": 7}
        mark_col = {"pass": 1, "fail": 3, "na": 4, "equip_na": 5}
        for key, r in row_index.items():
            entry = data["onsite_tests"].get(key, {})
            mark = entry.get("mark")
            if mark in mark_col:
                check_cell(t.cell(r, mark_col[mark]))
            if entry.get("remark"):
                set_cell_text(t.cell(r, 6), entry["remark"])

    if "shrinkage" in data:
        t = tables[TABLES["shrinkage_chart"]]
        # Delete existing rows except headers (0 and 1)
        for row in t.rows[2:]:
            t._tbl.remove(row._tr)
            
        s_data = data["shrinkage"]
        if isinstance(s_data, list):
            for item in s_data:
                if item.get("type") == "header":
                    r = t.add_row()
                    set_cell_text(r.cells[0], item.get("color", ""))
                    for i in range(1, min(3, len(r.cells))):
                        r.cells[0].merge(r.cells[i])
                elif item.get("type") == "data":
                    r = t.add_row()
                    set_cell_text(r.cells[0], str(item.get("before", "")))
                    set_cell_text(r.cells[1], str(item.get("after", "")))
                    set_cell_text(r.cells[2], str(item.get("pct", "")))

    if "marking_labeling" in data:
        t = tables[TABLES["marking_labeling"]]
        mark_col = {"conform": 1, "not_conform": 2, "na": 3}
        for row_str, entry in data["marking_labeling"].items():
            row = int(row_str)
            if row >= len(t.rows):
                continue
            mark = entry.get("mark")
            if mark in mark_col:
                check_cell(t.cell(row, mark_col[mark]))
            if entry.get("observation"):
                set_cell_text(t.cell(row, 4), entry["observation"])

    if "packing_matrix" in data:
        t = tables[TABLES["packing_matrix"]]
        for row_str, selected_indices in data["packing_matrix"].items():
            row = int(row_str)
            if row >= len(t.rows):
                continue
            all_cbs = []
            seen_cells = set()
            for cell in t.rows[row].cells:
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                boxes = _find_checkboxes_in_cell(cell)
                for cb, _ in boxes:
                    all_cbs.append(cb)
            
            for idx in selected_indices:
                if isinstance(idx, int) and 0 <= idx < len(all_cbs):
                    _set_checkbox_element(all_cbs[idx], True)
                    _insert_tick_after_checkbox(all_cbs[idx])

    if "photos" in data:
        for section, items in data["photos"].items():
            if section not in TABLES:
                continue
            t = tables[TABLES[section]]
            for item in items:
                row, col = item["row"], item["col"]
                img_cell = t.rows[row].cells[col]
                insert_image_in_cell(img_cell, item["path"])
                cap_row = row + 1
                if item.get("title") and cap_row < len(t.rows):
                    set_cell_text(t.rows[cap_row].cells[col], item["title"])

    doc.save(output_path)
    return output_path
